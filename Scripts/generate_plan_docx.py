"""Generate a professional Word document from the Rally Auto-Generation Plan."""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Add background color to a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_plan_document():
    """Create a professional Word document for the Rally Auto-Generation Plan."""

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Rally Test Case Auto-Generation Plan')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Voca Tooki Automation Framework')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Created: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()  # Spacing

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    summary_points = [
        'Vision: Rally becomes the single source of truth for all test cases',
        'Any test case marked as "automated" in Rally is automatically integrated into the framework',
        'Test structure synced to rally_suite.json AND pytest code auto-generated from test steps',
        'Zero manual test file creation — all pytest generated automatically from Rally data',
        'Current State: 2–3 manually coded test cases',
        'Target: Dozens of test cases auto-generated and runnable within 3 weeks',
        'Outcome: Scale from manual test case creation to fully automated generation'
    ]

    for point in summary_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Current State Table
    doc.add_heading('Current State', level=2)

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'Status'
    header_cells[2].text = 'Notes'

    for cell in header_cells:
        shade_cell(cell, '0033CC')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    data = [
        ['Manual TF/TC registration', 'Done', '3 folders, 2 cases in rally_suite.json'],
        ['Pytest auto-generation from Rally', 'Missing', 'Generated directly from test steps'],
        ['Rally API connection', 'Missing', 'Must be built'],
        ['Test sync from Rally', 'Missing', 'Must be built'],
        ['Pytest code generator', 'Missing', 'Must be built'],
        ['Automated flag filtering', 'Missing', 'Rally field not checked'],
        ['Full automation pipeline', 'Vision', 'This plan'],
        ['Timeline', 'Target', '3 weeks to completion'],
    ]

    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]

        # Color status cells
        if 'Missing' in row_data[1]:
            shade_cell(row_cells[1], 'FF9999')
        elif 'Done' in row_data[1]:
            shade_cell(row_cells[1], '99FF99')
        elif 'Target' in row_data[1] or 'Vision' in row_data[1]:
            shade_cell(row_cells[1], 'ADD8E6')

    doc.add_paragraph()

    # Phase 1
    doc.add_heading('Phase 1 — Rally Connection & Discovery', level=2)
    doc.add_paragraph('Week 1 | 2–3 days effort', style='Normal')

    doc.add_heading('Goal', level=3)
    doc.add_paragraph('Read test cases from Rally, identify which ones are marked "automated".', style='Normal')

    doc.add_heading('What Needs to Happen', level=3)

    phase1_tasks = [
        'Get Rally API Credentials (BLOCKING)',
        '  • Rally server URL (e.g., https://rally1.rallydev.com)',
        '  • API key or username/password',
        '  • Workspace ID and Project ID',
        '  • Confirm Rally has "Automated: true/false" field on TestCase',
        'Create runner/rally_sync.py — Rally API Client',
        '  • Authenticate to Rally WSAPI',
        '  • Query all TestCase objects filtered by Automated = true',
        '  • Extract: TC ID, name, folder, description, steps, expected results',
        '  • Parse test steps into structured format',
        'Create scripts/sync_rally_suite.py — Orchestration',
        '  • Call rally_sync.py to fetch all automated test cases',
        '  • Populate data/rally_suite.json with TF/TC hierarchy',
        '  • Log discovery results and any conflicts',
    ]

    for task in phase1_tasks:
        if task.startswith('  •'):
            p = doc.add_paragraph(task, style='List Bullet 2')
        elif task.startswith(' '):
            p = doc.add_paragraph(task[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('Deliverable', level=3)
    deliverables_p1 = [
        'runner/rally_sync.py — reads from Rally API',
        'scripts/sync_rally_suite.py — orchestrates sync',
        'Updated rally_suite.json with all automated TCs from Rally'
    ]
    for d in deliverables_p1:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()

    # Phase 2
    doc.add_heading('Phase 2 — 100% Auto-Generate Pytest (Zero Manual Coding)', level=2)
    doc.add_paragraph('Week 2 | 3–4 days effort', style='Normal')

    doc.add_heading('Goal', level=3)
    doc.add_paragraph('Convert Rally test case steps into complete, working pytest files — zero manual test coding required.', style='Normal')

    doc.add_heading('What Needs to Happen', level=3)

    phase2_tasks = [
        'Analyze Rally Test Case Structure',
        '  • Understand how test steps are formatted in Rally',
        '  • Extract: step actions, expected results, test data',
        '  • Map step actions to Page Object methods automatically',
        'Create runner/test_generator.py — Full Pytest Generator',
        '  • Takes Rally test case (ID, name, description, steps, results)',
        '  • Generates complete pytest function with all imports and fixtures',
        '  • Creates assertions directly from expected results',
        '  • Handles multi-step flows and data validation',
        '  • Output: ready-to-run pytest (no manual edits needed)',
        'Update scripts/sync_rally_suite.py',
        '  • Call test generator for each TC after syncing JSON',
        '  • Create Tests/rally/TF<n>_<name>/ directory structure',
        '  • Auto-generate test_tc<ID>_<name>.py files',
        '  • All generated files are immediately runnable',
    ]

    for task in phase2_tasks:
        if task.startswith('  •'):
            p = doc.add_paragraph(task, style='List Bullet 2')
        elif task.startswith(' '):
            p = doc.add_paragraph(task[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('Quality Gates', level=3)
    quality_checks = [
        'Generated pytest must be 100% syntactically valid',
        'Must use correct fixtures (altdriver) and Page Objects',
        'Must have complete assertions from expected results',
        'All imports must be present and correct',
        'Generated files must run immediately without manual edits',
        'Developer involvement only for new Page Object methods, not test coding'
    ]
    for check in quality_checks:
        doc.add_paragraph(check, style='List Bullet')

    doc.add_heading('Deliverable', level=3)
    deliverables_p2 = [
        'runner/test_generator.py — generates complete pytest from test steps',
        'Auto-generated pytest files in Tests/rally/ (all ready to run)',
        'Enhanced sync script with full code generation pipeline'
    ]
    for d in deliverables_p2:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()

    # Phase 3
    doc.add_heading('Phase 3 — Integration & Full Validation', level=2)
    doc.add_paragraph('Week 3 | 2–3 days effort', style='Normal')

    doc.add_heading('Goal', level=3)
    doc.add_paragraph('End-to-end: Rally → auto-generate → run all tests successfully.', style='Normal')

    doc.add_heading('What Needs to Happen', level=3)

    phase3_tasks = [
        'Wire Sync into Runner',
        '  • Add --sync-rally CLI flag to run_panel.py',
        '  • Or create "Sync from Rally" button in UI',
        '  • All tests auto-generated, zero manual work',
        'Configure rally.env',
        '  • RALLY_SERVER, RALLY_API_KEY, RALLY_WORKSPACE, RALLY_PROJECT',
        '  • Add to .gitignore (not committed)',
        'Validate Full Pipeline',
        '  • Run sync_rally_suite.py',
        '  • Verify all TCs have entries in rally_suite.json',
        '  • Verify all TCs have auto-generated pytest files',
        '  • Run all generated tests from Flask UI',
        '  • Confirm all tests pass without manual modifications',
    ]

    for task in phase3_tasks:
        if task.startswith('  •'):
            p = doc.add_paragraph(task, style='List Bullet 2')
        elif task.startswith(' '):
            p = doc.add_paragraph(task[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('Deliverable', level=3)
    deliverables_p3 = [
        'Full sync and generation working end-to-end',
        'All Rally automated TCs auto-generated and runnable',
        'Documentation for maintenance'
    ]
    for d in deliverables_p3:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()

    # Phase 4
    doc.add_heading('Phase 4 — Continuous Sync & Maintenance', level=2)
    doc.add_paragraph('Ongoing | 1–2 days setup', style='Normal')

    doc.add_heading('Goal', level=3)
    doc.add_paragraph('Keep Rally and framework in sync automatically.', style='Normal')

    doc.add_heading('What Needs to Happen', level=3)

    phase4_items = [
        'Automated Sync Hook',
        '  • Schedule sync to run daily or on-demand',
        '  • Or add "Sync Now" button to UI',
        'Change Detection',
        '  • New Rally TC → auto-generate pytest',
        '  • Deleted TC → remove from suite',
        '  • Modified TC → regenerate pytest',
        'Dev Workflow',
        '  • New Rally TC marked "automated" → next sync generates it',
        '  • No manual test file creation',
        '  • Devs only add Page Objects when needed',
    ]

    for item in phase4_items:
        if item.startswith('  •'):
            p = doc.add_paragraph(item, style='List Bullet 2')
        elif item.startswith(' '):
            p = doc.add_paragraph(item[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_page_break()

    # Timeline & Effort Table
    doc.add_heading('Timeline & Effort Summary', level=2)

    timeline_table = doc.add_table(rows=6, cols=4)
    timeline_table.style = 'Light Grid Accent 1'

    timeline_headers = timeline_table.rows[0].cells
    timeline_headers[0].text = 'Phase'
    timeline_headers[1].text = 'Duration'
    timeline_headers[2].text = 'Effort'
    timeline_headers[3].text = 'Outcome'

    for cell in timeline_headers:
        shade_cell(cell, '0033CC')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    timeline_data = [
        ['Phase 1 — Rally Discovery', '1 week', '2–3 days', 'Identify all automated TCs in Rally'],
        ['Phase 2 — 100% Pytest Auto-Generation', '1 week', '3–4 days', 'All TCs auto-generated, zero manual coding'],
        ['Phase 3 — E2E Validation', '1 week', '2–3 days', 'All tests running successfully'],
        ['Phase 4 — Continuous Sync', 'Ongoing', '1–2 days setup', 'New TCs auto-generated automatically'],
        ['TOTAL', '~3 weeks', '8–12 days', 'Full Rally automation with zero test coding'],
    ]

    for i, row_data in enumerate(timeline_data, 1):
        row_cells = timeline_table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        row_cells[3].text = row_data[3]

        if i == len(timeline_data):  # Last row
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                shade_cell(cell, 'E6E6FA')

    doc.add_paragraph()

    # Key Benefits
    doc.add_heading('Why This Approach Transforms Your Testing', level=2)

    benefits = [
        'Zero Manual Test Coding: All pytest generated automatically from Rally test steps',
        'Eliminate Manual Work: No developers writing test code, only maintaining Page Objects',
        'Faster Scaling: 100 new test cases auto-generated in hours instead of weeks',
        'Single Source of Truth: Rally is master, framework always in perfect sync',
        'Consistency: All generated tests follow identical patterns and structure',
        'Velocity Jump: From 2 hours per test to 5 minutes per test (full automation)',
        'Future-Proof: Every new Rally TC automatically becomes a runnable test'
    ]

    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_paragraph()

    # Additional Recommendations
    doc.add_heading('Additional Recommendations', level=2)

    recommendations = {
        'Rally Step Format': [
            'Ensure test cases have structured steps (JSON or consistent format)',
            'Each step: action, expected_result, optional test_data',
            'This enables reliable, high-fidelity auto-generation'
        ],
        'Page Object Inventory': [
            'Audit existing Pages: LoginPage, StartScreen, MapPage',
            'Document available methods',
            'Generator uses these for auto-mapping step actions'
        ],
        'Test Data Management': [
            'Pull from data/test_users.py or Rally TestData field',
            'Optional: parameterize generated tests for multiple data sets'
        ],
        'CI/CD Integration': [
            'Connect to GitHub Actions / Jenkins for auto-sync',
            'Auto-run all tests on each build',
            'Optional: report results back to Rally'
        ]
    }

    for title, items in recommendations.items():
        doc.add_heading(title, level=3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Success Metrics
    doc.add_heading('Success Metrics', level=2)

    metrics = {
        'Week 1': '15+ automated TCs discovered in Rally',
        'Week 2': '100% have auto-generated, ready-to-run pytest',
        'Week 3': 'All generated tests passing, zero manual edits needed',
        'Month 1': 'New Rally TC → auto-generated within 1 sync cycle',
        'Ongoing': '5 minutes per test (auto) vs. 2 hours (manual)'
    ]

    for period, metric in metrics.items():
        p = doc.add_paragraph()
        p.add_run(f'{period}: ').bold = True
        p.add_run(metric)

    doc.add_paragraph()

    # Final Recommendation
    doc.add_heading('Recommendation', level=2)

    rec_para = doc.add_paragraph(
        'This is a game-changing approach: it eliminates manual test code creation entirely. '
        'Every test case in Rally automatically becomes a runnable test in the framework.\n\n'
        'Start with Phase 1 immediately to unlock this productivity gain. '
        'The 3-week timeline is realistic and delivers full automation by end of Month 1.'
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f'---\nDocument prepared for management review\n{datetime.now().strftime("%B %d, %Y")}')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    return doc

if __name__ == '__main__':
    print('Generating Rally Auto-Generation Plan Word document...')
    doc = create_plan_document()
    output_path = os.path.join(
        os.path.dirname(__file__),
        '..',
        'Rally_Auto_Generation_Plan.docx'
    )
    doc.save(output_path)
    print('[OK] Document created successfully')
    print('Location: C:\\Users\\sanad\\PycharmProjects\\Automation\\Rally_Auto_Generation_Plan.docx')
    print('Ready to present to your manager!')
