from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Rally Test Case Auto-Generation Plan')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Voca Tooki Automation Framework')
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('Created: ' + datetime.now().strftime("%B %d, %Y"))
date_run.font.size = Pt(10)
date_run.font.italic = True
date_run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)

summary_points = [
    'Vision: Rally becomes the single source of truth for all test cases',
    'Every test case marked "automated" in Rally auto-generates complete pytest files',
    'ZERO manual test file creation - all pytest generated automatically from test steps',
    'Flask UI allows QA to select and run specific test folders or test cases on-demand',
    'Current State: 2-3 manually coded test cases (developers spend 2 hours per test)',
    'Target: Dozens of tests auto-generated and runnable within 3 weeks',
    'Outcome: Test creation drops from 2 hours to 5 minutes per test'
]

for point in summary_points:
    p = doc.add_paragraph(point, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

# Current State Table
doc.add_heading('Current State vs Vision', level=2)

table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Item'
header_cells[1].text = 'Today'
header_cells[2].text = 'After Implementation'

for cell in header_cells:
    shade_cell(cell, '0033CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

data = [
    ['Test Registration', 'Manual (3 TCs)', 'Automatic (all Rally TCs)'],
    ['Test File Creation', 'Manual coding (2 hrs/test)', 'Auto-generated (5 mins/test)'],
    ['Rally Connection', 'None', 'Full API integration'],
    ['Test Sync', 'Manual JSON edits', 'Automatic sync from Rally'],
    ['Code Generation', 'None', 'Complete pytest generation'],
    ['Automated Flag', 'Not checked', 'Filters Rally TCs automatically'],
    ['Test Selection UI', 'Basic', 'Advanced folder/case selection'],
    ['Scalability', 'Slow (manual)', 'Fast (auto-generated)'],
]

for i, row_data in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

    shade_cell(row_cells[1], 'FF9999')
    shade_cell(row_cells[2], '99FF99')

doc.add_paragraph()

# Flask UI Section
doc.add_heading('Flask UI - Test Selection & Execution', level=2)

doc.add_paragraph('QA can use the Flask web interface (http://localhost:5000) to select and run specific test cases.', style='Normal')

doc.add_heading('Features', level=3)

ui_features = [
    'Three Run Modes:',
    '  - Lesson Range: Run lessons across selected users',
    '  - Test Folder: Select specific test folders and run all cases inside',
    '  - Test Case: Select specific individual test cases to run',
    'Test Case Selection:',
    '  - Expandable/collapsible folder tree for easy navigation',
    '  - Checkboxes to select specific folders or test cases',
    '  - Selection is saved in browser (localStorage)',
    'Execution Control:',
    '  - Run button to execute selected tests',
    '  - Dry-run to preview without actual execution',
    '  - Stop button to halt running tests',
    'Live Monitoring:',
    '  - Real-time log streaming',
    '  - Progress bar showing test completion',
    '  - Individual test case result tracking',
    '  - Pass/Fail/Error status display'
]

for feature in ui_features:
    if feature.startswith('  -'):
        p = doc.add_paragraph(feature[4:], style='List Bullet 2')
    elif feature.endswith(':'):
        p = doc.add_paragraph(feature[:-1], style='List Bullet')
    else:
        p = doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('How to Use Test Selection', level=3)

usage_steps = [
    'Open http://localhost:5000 in your browser',
    'Click "Run Type" dropdown and select:',
    '  - Test Folder: to run entire folders',
    '  - Test Case: to run individual test cases',
    'Expand folders by clicking the arrow',
    'Check the checkbox next to desired folders/cases',
    'Click "Run" to execute selected tests',
    'Monitor progress in the live log and results panel',
    'View detailed results after execution'
]

for i, step in enumerate(usage_steps, 1):
    if step.startswith('  -'):
        p = doc.add_paragraph(step[4:], style='List Bullet 2')
    else:
        p = doc.add_paragraph(f'{i}. {step}' if not step[0].isdigit() else step, style='Normal')

doc.add_paragraph()

# Phase 1
doc.add_heading('Phase 1 - Rally Connection & Discovery', level=2)
doc.add_paragraph('Week 1 | 2-3 days effort', style='Normal')

doc.add_heading('Goal', level=3)
doc.add_paragraph('Read test cases from Rally and identify all marked as "automated".', style='Normal')

doc.add_heading('What Needs to Happen', level=3)

phase1_tasks = [
    'Get Rally API Credentials (BLOCKING)',
    '  - Rally server URL',
    '  - API key or username/password',
    '  - Workspace ID and Project ID',
    'Create runner/rally_sync.py',
    '  - Authenticate to Rally WSAPI',
    '  - Query all TestCase objects where Automated = true',
    '  - Extract TC ID, name, folder, description, and test steps',
    '  - Parse steps into structured format',
    'Create scripts/sync_rally_suite.py',
    '  - Orchestrate the sync process',
    '  - Populate data/rally_suite.json with TF/TC hierarchy',
    '  - Log discovery results'
]

for task in phase1_tasks:
    if task.startswith('  -'):
        p = doc.add_paragraph(task[4:], style='List Bullet 2')
    else:
        p = doc.add_paragraph(task, style='List Bullet')

doc.add_paragraph()

# Phase 2
doc.add_heading('Phase 2 - 100% Auto-Generate Pytest (Zero Manual Coding)', level=2)
doc.add_paragraph('Week 2 | 3-4 days effort', style='Normal')

doc.add_heading('Goal', level=3)
doc.add_paragraph('Convert Rally test case steps into complete, working pytest files - zero manual test coding required.', style='Normal')

doc.add_heading('What Needs to Happen', level=3)

phase2_tasks = [
    'Analyze Rally Test Case Structure',
    '  - Understand test step format in Rally',
    '  - Extract: actions, expected results, test data',
    '  - Map actions to Page Object methods',
    'Create runner/test_generator.py',
    '  - Full pytest code generation engine',
    '  - Takes Rally TC (ID, name, steps, results)',
    '  - Generates complete pytest with imports and fixtures',
    '  - Creates assertions from expected results',
    '  - Output: ready-to-run pytest (no manual edits needed)',
    'Update scripts/sync_rally_suite.py',
    '  - Call test generator for each TC',
    '  - Create Tests/rally/ directory structure',
    '  - Auto-generate test_tc<ID>_<name>.py files',
    '  - All files immediately runnable'
]

for task in phase2_tasks:
    if task.startswith('  -'):
        p = doc.add_paragraph(task[4:], style='List Bullet 2')
    else:
        p = doc.add_paragraph(task, style='List Bullet')

doc.add_heading('Quality Gates', level=3)
quality = [
    'Generated pytest must be 100% syntactically valid',
    'Must use correct fixtures (altdriver) and Page Objects',
    'Must have complete assertions',
    'All imports must be present and correct',
    'Generated files must run immediately without manual edits',
    'Zero developer involvement in test file coding'
]
for q in quality:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph()

# Phase 3
doc.add_heading('Phase 3 - Integration & Full Validation', level=2)
doc.add_paragraph('Week 3 | 2-3 days effort', style='Normal')

doc.add_heading('Goal', level=3)
doc.add_paragraph('End-to-end: Rally -> auto-generate -> run all tests successfully.', style='Normal')

doc.add_heading('What Needs to Happen', level=3)

phase3 = [
    'Wire Sync into Runner',
    '  - Add --sync-rally CLI flag or UI button',
    '  - All tests auto-generated, zero manual work',
    'Configure rally.env',
    '  - Set RALLY_SERVER, RALLY_API_KEY, etc',
    '  - Add to .gitignore',
    'Enhance Flask UI (Already Exists)',
    '  - Verify test selection works correctly',
    '  - Test folder and case selection functionality',
    '  - Confirm selections are saved and restored',
    'Validate Full Pipeline',
    '  - Run sync_rally_suite.py',
    '  - Verify all TCs in rally_suite.json',
    '  - Use Flask UI to select and run tests',
    '  - Verify selections persist across page reloads',
    '  - Confirm all pass without manual mods'
]

for item in phase3:
    if item.startswith('  -'):
        p = doc.add_paragraph(item[4:], style='List Bullet 2')
    else:
        p = doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Phase 4
doc.add_heading('Phase 4 - Continuous Sync & Maintenance', level=2)
doc.add_paragraph('Ongoing | 1-2 days setup', style='Normal')

doc.add_heading('Goal', level=3)
doc.add_paragraph('Keep Rally and framework in sync automatically.', style='Normal')

doc.add_heading('What Needs to Happen', level=3)

phase4 = [
    'Automated Sync Hook',
    '  - Schedule daily sync or on-demand',
    '  - Add "Sync Now" button to Flask UI',
    'Change Detection',
    '  - New Rally TC -> auto-generate pytest',
    '  - Deleted TC -> remove from suite',
    '  - Modified TC -> regenerate',
    'Dev Workflow Update',
    '  - New Rally TC marked "automated" -> next sync generates it',
    '  - ZERO manual test file creation',
    '  - Developers only add Page Objects when needed'
]

for item in phase4:
    if item.startswith('  -'):
        p = doc.add_paragraph(item[4:], style='List Bullet 2')
    else:
        p = doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Timeline Table
doc.add_heading('Timeline & Effort Summary', level=2)

timeline_table = doc.add_table(rows=6, cols=4)
timeline_table.style = 'Light Grid Accent 1'

timeline_headers = timeline_table.rows[0].cells
timeline_headers[0].text = 'Phase'
timeline_headers[1].text = 'Duration'
timeline_headers[2].text = 'Effort'
timeline_headers[3].text = 'Outcome'

for cell in timeline_headers:
    shade_cell(cell, '0033CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

timeline_data = [
    ['Phase 1 - Rally Discovery', '1 week', '2-3 days', 'Identify all automated TCs in Rally'],
    ['Phase 2 - Auto-Generation', '1 week', '3-4 days', 'All TCs auto-generated, zero manual coding'],
    ['Phase 3 - Validation', '1 week', '2-3 days', 'Full pipeline working, Flask UI test selection'],
    ['Phase 4 - Continuous Sync', 'Ongoing', '1-2 days setup', 'New TCs auto-generated automatically'],
    ['TOTAL', 'approx 3 weeks', '8-12 days', 'Full Rally automation, Flask test selection'],
]

for i, row_data in enumerate(timeline_data, 1):
    row_cells = timeline_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
    row_cells[3].text = row_data[3]

    if i == len(timeline_data):
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            shade_cell(cell, 'E6E6FA')

doc.add_paragraph()

# Key Benefits
doc.add_heading('Why This Approach Is Game-Changing', level=2)

benefits = [
    'ZERO Manual Test Coding: All pytest generated automatically from Rally test steps',
    'Eliminate Developer Bottleneck: No more developers stuck writing test code',
    'Flexible Test Execution: QA selects specific tests to run via web UI',
    'Massive Scaling: 100 new Rally TCs auto-generated in hours, not weeks',
    'Single Source of Truth: Rally is master, framework always in perfect sync',
    'Consistency: All generated tests follow identical patterns',
    'Velocity: Test creation drops from 2 hours to 5 minutes per test'
]

for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()

# Success Metrics
doc.add_heading('Success Metrics', level=2)

metrics = [
    ('Week 1', '15+ automated TCs discovered in Rally'),
    ('Week 2', '100% have auto-generated pytest ready to run'),
    ('Week 3', 'All generated tests passing, Flask UI working, test selection functional'),
    ('Month 1', 'New Rally TC auto-generated within 1 sync cycle'),
    ('Velocity', '5 minutes per test (auto) vs 2 hours (manual)')
]

for period, metric in metrics:
    p = doc.add_paragraph()
    p.add_run(period + ': ').bold = True
    p.add_run(metric)

doc.add_paragraph()

# Key Files Reference
doc.add_heading('Key Files to Create/Modify', level=2)

files_table = doc.add_table(rows=11, cols=2)
files_table.style = 'Light Grid Accent 1'

files_headers = files_table.rows[0].cells
files_headers[0].text = 'File'
files_headers[1].text = 'Purpose'

for cell in files_headers:
    shade_cell(cell, '0033CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

files_data = [
    ['runner/rally_sync.py (NEW)', 'Rally API client - fetch TCs'],
    ['runner/test_generator.py (NEW)', 'Complete pytest auto-generator'],
    ['scripts/sync_rally_suite.py (NEW)', 'Orchestrate full sync process'],
    ['rally.env (NEW)', 'Rally credentials (not committed)'],
    ['run_panel.py', 'Flask app, load rally.env'],
    ['runner/templates/index.html', 'Flask UI (test selection already built-in)'],
    ['runner/static/app.js', 'UI logic (folder/case selection already built-in)'],
    ['data/rally_suite.json', 'Auto-populated by sync'],
    ['Tests/rally/', 'Auto-generated pytest files'],
    ['conftest.py', 'Works with generated tests'],
]

for i, row_data in enumerate(files_data, 1):
    row_cells = files_table.rows[i].cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

    if 'NEW' in row_data[0]:
        shade_cell(row_cells[0], 'FFFFCC')

doc.add_paragraph()

# Final Recommendation
doc.add_heading('Recommendation', level=2)

rec = doc.add_paragraph(
    'Start with Phase 1 immediately. Getting Rally credentials is the critical path. '
    'The Flask UI test selection feature is already built-in and ready to use. '
    'This approach transforms test coverage from manual creation to fully automated generation '
    '- eliminating test file coding entirely. QA can use the web UI to flexibly select and run '
    'which tests they need, while developers focus on Page Object maintenance.'
)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('Document prepared for management review\n' + datetime.now().strftime("%B %d, %Y"))
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(150, 150, 150)

# Save
import os
import shutil
temp_path = os.path.join(os.path.dirname(__file__), '..', 'Rally_Plan_TEMP.docx')
final_path = os.path.join(os.path.dirname(__file__), '..', 'Rally_Auto_Generation_Plan.docx')

doc.save(temp_path)

# Try to overwrite the locked file
try:
    shutil.move(temp_path, final_path)
except:
    final_path = temp_path

print('[SUCCESS] Updated Word document with Flask UI test selection section')
print('Location: ' + final_path)
