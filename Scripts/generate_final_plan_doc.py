from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

doc = Document()

sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Automation Framework')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Must-Have Improvements & Rally Integration Plan')
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2_run = subtitle2.add_run('Voca Tooki — 14-18 Days to Production Grade')
subtitle2_run.font.size = Pt(12)
subtitle2_run.font.italic = True

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('Created: ' + datetime.now().strftime("%B %d, %Y"))
date_run.font.size = Pt(10)
date_run.font.italic = True
date_run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)

summary = [
    'CRITICAL: Plaintext passwords are committed to the repository in 5 locations',
    'Framework is functional but lacks resilience features (retries, timeouts, screenshots)',
    'No CI/CD pipeline — tests only run manually via .bat files',
    'Rally integration planned in 4 phases over 3 weeks',
    'Must-have improvements must be fixed alongside Rally integration',
    'Total effort: 14-18 days to production-grade automation'
]

for s in summary:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()

# Gap Analysis Table
doc.add_heading('Critical Gaps Analysis', level=2)

table = doc.add_table(rows=15, cols=3)
table.style = 'Light Grid Accent 1'

headers = table.rows[0].cells
headers[0].text = 'Area'
headers[1].text = 'Status'
headers[2].text = 'Priority'

for cell in headers:
    shade_cell(cell, '0033CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

gaps = [
    ('Credential security', 'Plaintext passwords in source files', 'P0 CRITICAL'),
    ('Screenshot on failure', 'Config exists, code does not', 'P1'),
    ('Per-test timeout', 'Tests can hang indefinitely', 'P1'),
    ('CI/CD pipeline', 'Only manual .bat files', 'P1'),
    ('Rally auto-sync + generation', 'Not implemented', 'P1'),
    ('Retry logic', 'No flaky test handling', 'P2'),
    ('Slack/Teams alerts', 'Email only', 'P2'),
    ('Environment config', 'Hardcoded to 127.0.0.1', 'P2'),
    ('Parallel execution', 'xdist installed but not used', 'P2'),
    ('Test tagging', 'Inconsistent markers', 'P3'),
    ('HTML reports', 'Basic, no charts/screenshots', 'P3'),
    ('Coverage metrics', 'Not tracked', 'P3'),
    ('Flask UI', 'Test selection already works', 'DONE'),
    ('Live logging', 'Real-time SSE streaming', 'DONE'),
]

for i, (area, status, priority) in enumerate(gaps, 1):
    row = table.rows[i].cells
    row[0].text = area
    row[1].text = status
    row[2].text = priority

    if 'CRITICAL' in priority:
        shade_cell(row[2], 'FF0000')
        shade_cell(row[1], 'FF9999')
    elif 'P1' in priority:
        shade_cell(row[2], 'FF9999')
    elif 'DONE' in priority:
        shade_cell(row[2], '99FF99')
        shade_cell(row[1], '99FF99')

doc.add_paragraph()

# P0 Section
doc.add_heading('P0 — CRITICAL: Credential Security (1 Day)', level=2)

doc.add_paragraph('Passwords found in:', style='Normal')
creds = [
    'data/test_users.py — "7173"',
    'Tests/rally/*/test_tc128_*.py — "0391"',
    'Tests/rally/*/test_tc132_*.py — hardcoded',
    'data/rally_suite.json — inline passwords',
    'automation_email.env — Gmail password (CRITICAL)',
]
for c in creds:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Action:', style='Normal')
fix = [
    'Move all credentials to .env files',
    'Add .env files to .gitignore',
    'Load from environment variables at runtime',
    'Files: data/test_users.py, conftest.py, Tests/rally/, run_panel.py'
]
for f in fix:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()

# P1 Improvements
doc.add_heading('P1 — Must Have Improvements (3 Days Total)', level=2)

improvements = {
    'Screenshot on Failure (0.5 day)': [
        'Add pytest_runtest_makereport hook to conftest.py',
        'Capture screenshot on test failure via AltTester API',
        'Store in screenshots/ folder for evidence'
    ],
    'Per-Test Timeout (0.5 day)': [
        'Add pytest-timeout to requirements.txt',
        'Set timeout = 120 in pytest.ini',
        'Prevents infinite hangs'
    ],
    'CI/CD Pipeline (1-2 days)': [
        'Create .github/workflows/automation.yml',
        'Trigger on schedule or game build',
        'Auto-run tests, upload results to Rally'
    ]
}

for title, items in improvements.items():
    doc.add_heading(title, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Rally Integration
doc.add_heading('Rally Integration Plan (3 Weeks)', level=2)

phases = [
    ('Phase 1: Rally Discovery (Week 1, 2-3 days)', [
        'Create runner/rally_sync.py — Rally API client',
        'Query all TestCases where Automated = true',
        'Populate data/rally_suite.json from Rally'
    ]),
    ('Phase 2: Auto-Generate Pytest (Week 2, 3-4 days)', [
        'Create runner/test_generator.py',
        'Convert Rally test steps to complete pytest files',
        'Zero manual test coding'
    ]),
    ('Phase 3: Flask UI Wiring (Week 3, 2-3 days)', [
        'Add "Sync from Rally" button to Flask UI',
        'Wire selected tests to run configuration',
        'Full end-to-end validation'
    ]),
    ('Phase 4: Continuous Sync (Ongoing, 1-2 days)', [
        'Auto-detect new/modified/deleted TCs',
        'Auto-generate pytest for new TCs',
        'Keep framework in sync with Rally'
    ])
]

for phase, details in phases:
    doc.add_heading(phase, level=3)
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')

doc.add_paragraph()

# Timeline
doc.add_heading('Full Timeline & Effort', level=2)

timeline_table = doc.add_table(rows=13, cols=4)
timeline_table.style = 'Light Grid Accent 1'

t_headers = timeline_table.rows[0].cells
t_headers[0].text = 'Task'
t_headers[1].text = 'Priority'
t_headers[2].text = 'Effort'
t_headers[3].text = 'Outcome'

for cell in t_headers:
    shade_cell(cell, '0033CC')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

timeline_items = [
    ('Credential security', 'P0', '1 day', 'Passwords out of source'),
    ('Screenshot on failure', 'P1', '0.5 day', 'Evidence for all failures'),
    ('Per-test timeout', 'P1', '0.5 day', 'No hanging tests'),
    ('CI/CD pipeline', 'P1', '1-2 days', 'Automated test execution'),
    ('Rally Phase 1', 'P1', '2-3 days', 'Discover all Rally TCs'),
    ('Rally Phase 2', 'P1', '3-4 days', 'Auto-generate all pytest'),
    ('Rally Phase 3', 'P1', '2-3 days', 'Full pipeline validation'),
    ('Retry flaky tests', 'P2', '0.5 day', 'Resilient testing'),
    ('Slack/Teams alerts', 'P2', '0.5 day', 'Instant team notifications'),
    ('Environment config', 'P2', '1 day', 'Dev/staging/prod support'),
    ('Rally Phase 4', 'P2 (ongoing)', '1-2 days', 'Continuous sync'),
    ('TOTAL', 'Mixed', '14-18 days', 'Production-grade framework'),
]

for i, (task, priority, effort, outcome) in enumerate(timeline_items, 1):
    row = timeline_table.rows[i].cells
    row[0].text = task
    row[1].text = priority
    row[2].text = effort
    row[3].text = outcome

    if i == len(timeline_items):
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            shade_cell(cell, 'E6E6FA')

doc.add_paragraph()

# Key Files
doc.add_heading('Key Files to Create/Modify', level=2)

files = [
    ('conftest.py', 'Add screenshot hook, timeout'),
    ('requirements.txt', 'Add pytest-timeout, pytest-rerunfailures'),
    ('pytest.ini', 'Add timeout, reruns config'),
    ('data/test_users.py', 'Load credentials from env'),
    ('runner/emailer.py', 'Add Slack webhook'),
    ('runner/core.py', 'Rally upload hook, Slack call'),
    ('run_panel.py', 'Load rally.env'),
    ('runner/templates/index.html', 'Add Sync button'),
    ('NEW: runner/rally_sync.py', 'Rally API client'),
    ('NEW: runner/test_generator.py', 'Pytest auto-generator'),
    ('NEW: scripts/sync_rally_suite.py', 'Orchestrate sync'),
    ('NEW: .github/workflows/automation.yml', 'CI/CD pipeline'),
]

for fname, role in files:
    p = doc.add_paragraph()
    if 'NEW:' in fname:
        p.add_run(fname + ' — ').bold = True
        p.add_run(role)
    else:
        p.add_run(fname + ' — ').italic = True
        p.add_run(role)

doc.add_paragraph()

# Recommendation
doc.add_heading('Recommendation', level=2)

rec = doc.add_paragraph(
    'IMMEDIATE ACTION: Fix credential security (P0) within 1 day. This is a security incident. '
    '\n\n'
    'PARALLEL TRACK: Begin Rally integration (Phase 1) immediately after or alongside security fix. '
    '\n\n'
    'TIMELINE: 14-18 days total to production-grade framework with full Rally integration, CI/CD, '
    'resilience features, and monitoring.'
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('Document prepared for management review\n' + datetime.now().strftime("%B %d, %Y"))
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(150, 150, 150)

import os
output = os.path.join(os.path.dirname(__file__), '..', 'Framework_Improvement_Plan.docx')
doc.save(output)
print('[OK] Professional plan document created')
print('File: Framework_Improvement_Plan.docx')
